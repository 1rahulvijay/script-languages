import docx
import requests
import re
import os
import matplotlib.pyplot as plt
from docx.shared import Inches


# TASK 2
def extractBookTitle(bookText):
    titleBeg = 'Title: '
    titleLine = re.search(titleBeg, bookText)
    return bookText[titleLine.span()[1]: bookText.find('\n', titleLine.span()[0])]


def extractAuthor(bookText):
    authBeg = 'Author: '
    author = re.search(authBeg, bookText)
    return bookText[author.span()[1]: bookText.find('\n', author.span()[0])]


def getChapter(bookText, beginning, end):
    try:
        startInd = bookText.index(beginning) + len(beginning)
        endInd = bookText.index(end, startInd)

        return bookText[startInd:endInd]
    except ValueError:
        return ""


# TASK 3
def rounded(num):
    return num - (num % 10)


def getParagraphArray(bookText):
    ParagraphsToCut = bookText.replace('\r', '').split('\n\n')
    ParagraphsSplit = []
    for string in ParagraphsToCut:
        if string != "" and string != '\n':
            ParagraphsSplit.append(string)
    return ParagraphsSplit


def getWordCountForParagraphs(paragraphs):
    CountedWordsInParagraphs = []
    for paragraph in paragraphs:
        CountedWordsInParagraphs.append(rounded(len(paragraph.replace('\n', " ").split(" "))))
    return sorted(CountedWordsInParagraphs)


# TASK 7
def projectReport(bookLink):
    print("Project Part 1 run successfully, Please wait")
    book = requests.get(bookLink)
    bookText = book.text.strip()

    chapterOne = getChapter(bookText, 'CHAPTER I.', 'CHAPTER II.')
    paragraphs = getParagraphArray(chapterOne)
    mps = getWordCountForParagraphs(paragraphs)

    # Creating plot
    print("Once this distribution has plotted close pyplot window to see Project report")
    plt.figure(figsize=(7, 5))
    plt.title("Words Distribution in PieChart")
    plt.pie(mps)
    plt.savefig('WordDistribution.png')

    # show plot
    plt.show()

    doc = docx.Document()

    doc.add_heading(extractBookTitle(bookText), 0)
    doc.add_heading('Author :' + extractAuthor(bookText))
    doc.add_heading('Report by: Rahul Vijayvargiya \n')
    doc.add_heading('Project Report submitted to Prof : Dr inż. Ryszard Świerczyński \n')

    r = doc.add_paragraph()

    r.add_run('Chapter I has: ' + str(len(paragraphs)) + 'paragraphs. \n').bold = True
    r.add_run('It also has: ' + str(len(chapterOne.split())) + ' words.\n').italic = True
    r.add_run('The average is: ' + str(round(sum(mps) / len(mps))) + ' words per paragraph \n').italic = True
    r.add_run('The longest paragraph has:  ' + str(max(mps)) + ' words \n').italic = True
    r.add_run('The shortest paragraph has:  ' + str(min(mps)) + ' words \n\n').italic = True
    r.add_run('The book is linked here:  ' + bookLink)
    doc.add_picture('WordDistribution.png', width=Inches(4.5))
    r.add_run('\n\n Thank You, Have a nice day').bold = True
    doc.save('projectReport.docx')
    os.system('start projectReport.docx')


def run():
    # TASK 1
    bookLink = 'https://www.gutenberg.org//cache/epub/4363/pg4363.txt'

    projectReport(bookLink)


if __name__ == "__main__":
    run()
